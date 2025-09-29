"""
预览服务
"""
import os
import asyncio
import hashlib
from pathlib import Path
from typing import Dict, List, Optional, Any
from PIL import Image, ImageOps
import logging
from app.database import SessionLocal
from app.models import Blob, Asset
from app.services.job_service import JobService

logger = logging.getLogger(__name__)

class PreviewService:
    """预览服务"""
    
    def __init__(self, cache_dir: str = "./cache/preview"):
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        self.job_service = JobService()
        
        # 支持的预览类型
        self.supported_types = {
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
            'document': ['.pdf', '.doc', '.docx', '.ppt', '.pptx', '.xls', '.xlsx'],
            'audio': ['.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a'],
            'video': ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm']
        }
        
        # 预览尺寸配置
        self.preview_sizes = {
            'small': (64, 64),
            'medium': (256, 256),
            'large': (512, 512),
            'thumbnail': (150, 150)
        }
    
    def get_preview_type(self, file_path: str) -> Optional[str]:
        """获取文件预览类型"""
        ext = Path(file_path).suffix.lower()
        
        for preview_type, extensions in self.supported_types.items():
            if ext in extensions:
                return preview_type
        
        return None
    
    def get_preview_path(self, content_hash: str, size: str = 'medium', preview_type: str = None) -> Path:
        """获取预览文件路径"""
        if preview_type:
            return self.cache_dir / content_hash / f"{size}.{preview_type}"
        else:
            return self.cache_dir / content_hash / f"{size}.jpg"
    
    def is_preview_cached(self, content_hash: str, size: str = 'medium') -> bool:
        """检查预览是否已缓存"""
        preview_path = self.get_preview_path(content_hash, size)
        return preview_path.exists()
    
    async def generate_preview(self, content_hash: str, file_path: str, size: str = 'medium') -> Dict[str, Any]:
        """生成文件预览"""
        try:
            # 检查是否已缓存
            if self.is_preview_cached(content_hash, size):
                return {
                    'success': True,
                    'cached': True,
                    'preview_path': str(self.get_preview_path(content_hash, size)),
                    'content_hash': content_hash
                }
            
            # 获取文件类型
            preview_type = self.get_preview_type(file_path)
            if not preview_type:
                return {
                    'success': False,
                    'error': '不支持的文件类型',
                    'content_hash': content_hash
                }
            
            # 创建预览目录
            preview_dir = self.cache_dir / content_hash
            preview_dir.mkdir(parents=True, exist_ok=True)
            
            # 根据类型生成预览
            if preview_type == 'image':
                result = await self.generate_image_preview(file_path, content_hash, size)
            elif preview_type == 'document':
                result = await self.generate_document_preview(file_path, content_hash, size)
            elif preview_type == 'audio':
                result = await self.generate_audio_preview(file_path, content_hash, size)
            elif preview_type == 'video':
                result = await self.generate_video_preview(file_path, content_hash, size)
            else:
                result = {
                    'success': False,
                    'error': '未知的预览类型',
                    'content_hash': content_hash
                }
            
            return result
            
        except Exception as e:
            logger.error(f"生成预览失败: {content_hash}, 错误: {e}")
            return {
                'success': False,
                'error': str(e),
                'content_hash': content_hash
            }
    
    async def generate_image_preview(self, file_path: str, content_hash: str, size: str) -> Dict[str, Any]:
        """生成图片预览"""
        try:
            # 打开图片
            with Image.open(file_path) as img:
                # 转换为RGB模式
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # 获取目标尺寸
                target_size = self.preview_sizes.get(size, (256, 256))
                
                # 保持宽高比缩放
                img.thumbnail(target_size, Image.Resampling.LANCZOS)
                
                # 如果图片小于目标尺寸，居中放置
                if img.size < target_size:
                    new_img = Image.new('RGB', target_size, (255, 255, 255))
                    paste_x = (target_size[0] - img.size[0]) // 2
                    paste_y = (target_size[1] - img.size[1]) // 2
                    new_img.paste(img, (paste_x, paste_y))
                    img = new_img
                
                # 保存预览
                preview_path = self.get_preview_path(content_hash, size, 'jpg')
                img.save(preview_path, 'JPEG', quality=85, optimize=True)
                
                return {
                    'success': True,
                    'preview_path': str(preview_path),
                    'preview_type': 'image',
                    'size': img.size,
                    'content_hash': content_hash
                }
                
        except Exception as e:
            logger.error(f"生成图片预览失败: {file_path}, 错误: {e}")
            return {
                'success': False,
                'error': f"图片预览生成失败: {str(e)}",
                'content_hash': content_hash
            }
    
    async def generate_document_preview(self, file_path: str, content_hash: str, size: str) -> Dict[str, Any]:
        """生成文档预览"""
        try:
            ext = Path(file_path).suffix.lower()
            
            if ext == '.pdf':
                return await self.generate_pdf_preview(file_path, content_hash, size)
            elif ext in ['.doc', '.docx']:
                return await self.generate_word_preview(file_path, content_hash, size)
            elif ext in ['.ppt', '.pptx']:
                return await self.generate_powerpoint_preview(file_path, content_hash, size)
            elif ext in ['.xls', '.xlsx']:
                return await self.generate_excel_preview(file_path, content_hash, size)
            else:
                return {
                    'success': False,
                    'error': f'不支持的文档类型: {ext}',
                    'content_hash': content_hash
                }
                
        except Exception as e:
            logger.error(f"生成文档预览失败: {file_path}, 错误: {e}")
            return {
                'success': False,
                'error': f"文档预览生成失败: {str(e)}",
                'content_hash': content_hash
            }
    
    async def generate_pdf_preview(self, file_path: str, content_hash: str, size: str) -> Dict[str, Any]:
        """生成PDF预览"""
        try:
            # 使用pdf2image生成PDF第一页预览
            from pdf2image import convert_from_path
            
            # 转换PDF第一页为图片
            pages = convert_from_path(file_path, first_page=1, last_page=1, dpi=150)
            
            if pages:
                img = pages[0]
                
                # 获取目标尺寸
                target_size = self.preview_sizes.get(size, (256, 256))
                
                # 保持宽高比缩放
                img.thumbnail(target_size, Image.Resampling.LANCZOS)
                
                # 保存预览
                preview_path = self.get_preview_path(content_hash, size, 'jpg')
                img.save(preview_path, 'JPEG', quality=85, optimize=True)
                
                return {
                    'success': True,
                    'preview_path': str(preview_path),
                    'preview_type': 'document',
                    'size': img.size,
                    'content_hash': content_hash
                }
            else:
                return {
                    'success': False,
                    'error': 'PDF文件为空或无法读取',
                    'content_hash': content_hash
                }
                
        except Exception as e:
            logger.error(f"生成PDF预览失败: {file_path}, 错误: {e}")
            return {
                'success': False,
                'error': f"PDF预览生成失败: {str(e)}",
                'content_hash': content_hash
            }
    
    async def generate_word_preview(self, file_path: str, content_hash: str, size: str) -> Dict[str, Any]:
        """生成Word文档预览"""
        try:
            # 使用python-docx读取文档
            from docx import Document
            
            doc = Document(file_path)
            
            # 创建预览图片
            target_size = self.preview_sizes.get(size, (256, 256))
            img = Image.new('RGB', target_size, (255, 255, 255))
            
            # 在图片上绘制文档信息
            from PIL import ImageDraw, ImageFont
            
            draw = ImageDraw.Draw(img)
            
            # 使用默认字体
            try:
                font = ImageFont.truetype("arial.ttf", 16)
            except:
                font = ImageFont.load_default()
            
            # 绘制文档标题
            title = doc.paragraphs[0].text if doc.paragraphs else "Word文档"
            draw.text((10, 10), title[:30], fill=(0, 0, 0), font=font)
            
            # 绘制段落数量
            draw.text((10, 40), f"段落数: {len(doc.paragraphs)}", fill=(100, 100, 100), font=font)
            
            # 保存预览
            preview_path = self.get_preview_path(content_hash, size, 'jpg')
            img.save(preview_path, 'JPEG', quality=85, optimize=True)
            
            return {
                'success': True,
                'preview_path': str(preview_path),
                'preview_type': 'document',
                'size': img.size,
                'content_hash': content_hash
            }
            
        except Exception as e:
            logger.error(f"生成Word预览失败: {file_path}, 错误: {e}")
            return {
                'success': False,
                'error': f"Word预览生成失败: {str(e)}",
                'content_hash': content_hash
            }
    
    async def generate_powerpoint_preview(self, file_path: str, content_hash: str, size: str) -> Dict[str, Any]:
        """生成PowerPoint预览"""
        try:
            # 创建预览图片
            target_size = self.preview_sizes.get(size, (256, 256))
            img = Image.new('RGB', target_size, (255, 255, 255))
            
            from PIL import ImageDraw, ImageFont
            
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("arial.ttf", 16)
            except:
                font = ImageFont.load_default()
            
            # 绘制PPT信息
            draw.text((10, 10), "PowerPoint演示文稿", fill=(0, 0, 0), font=font)
            draw.text((10, 40), "点击查看完整内容", fill=(100, 100, 100), font=font)
            
            # 保存预览
            preview_path = self.get_preview_path(content_hash, size, 'jpg')
            img.save(preview_path, 'JPEG', quality=85, optimize=True)
            
            return {
                'success': True,
                'preview_path': str(preview_path),
                'preview_type': 'document',
                'size': img.size,
                'content_hash': content_hash
            }
            
        except Exception as e:
            logger.error(f"生成PowerPoint预览失败: {file_path}, 错误: {e}")
            return {
                'success': False,
                'error': f"PowerPoint预览生成失败: {str(e)}",
                'content_hash': content_hash
            }
    
    async def generate_excel_preview(self, file_path: str, content_hash: str, size: str) -> Dict[str, Any]:
        """生成Excel预览"""
        try:
            # 创建预览图片
            target_size = self.preview_sizes.get(size, (256, 256))
            img = Image.new('RGB', target_size, (255, 255, 255))
            
            from PIL import ImageDraw, ImageFont
            
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("arial.ttf", 16)
            except:
                font = ImageFont.load_default()
            
            # 绘制Excel信息
            draw.text((10, 10), "Excel电子表格", fill=(0, 0, 0), font=font)
            draw.text((10, 40), "点击查看完整内容", fill=(100, 100, 100), font=font)
            
            # 保存预览
            preview_path = self.get_preview_path(content_hash, size, 'jpg')
            img.save(preview_path, 'JPEG', quality=85, optimize=True)
            
            return {
                'success': True,
                'preview_path': str(preview_path),
                'preview_type': 'document',
                'size': img.size,
                'content_hash': content_hash
            }
            
        except Exception as e:
            logger.error(f"生成Excel预览失败: {file_path}, 错误: {e}")
            return {
                'success': False,
                'error': f"Excel预览生成失败: {str(e)}",
                'content_hash': content_hash
            }
    
    async def generate_audio_preview(self, file_path: str, content_hash: str, size: str) -> Dict[str, Any]:
        """生成音频预览"""
        try:
            # 创建音频波形预览
            target_size = self.preview_sizes.get(size, (256, 256))
            img = Image.new('RGB', target_size, (240, 240, 240))
            
            from PIL import ImageDraw, ImageFont
            
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("arial.ttf", 16)
            except:
                font = ImageFont.load_default()
            
            # 绘制音频信息
            draw.text((10, 10), "音频文件", fill=(0, 0, 0), font=font)
            draw.text((10, 40), "点击播放", fill=(100, 100, 100), font=font)
            
            # 绘制简单的波形图
            width, height = target_size
            for i in range(0, width-20, 4):
                y = height // 2 + (i % 20 - 10)
                draw.line([(i+10, height//2), (i+10, y)], fill=(0, 100, 200), width=2)
            
            # 保存预览
            preview_path = self.get_preview_path(content_hash, size, 'jpg')
            img.save(preview_path, 'JPEG', quality=85, optimize=True)
            
            return {
                'success': True,
                'preview_path': str(preview_path),
                'preview_type': 'audio',
                'size': img.size,
                'content_hash': content_hash
            }
            
        except Exception as e:
            logger.error(f"生成音频预览失败: {file_path}, 错误: {e}")
            return {
                'success': False,
                'error': f"音频预览生成失败: {str(e)}",
                'content_hash': content_hash
            }
    
    async def generate_video_preview(self, file_path: str, content_hash: str, size: str) -> Dict[str, Any]:
        """生成视频预览"""
        try:
            # 创建视频预览
            target_size = self.preview_sizes.get(size, (256, 256))
            img = Image.new('RGB', target_size, (20, 20, 20))
            
            from PIL import ImageDraw, ImageFont
            
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("arial.ttf", 16)
            except:
                font = ImageFont.load_default()
            
            # 绘制视频信息
            draw.text((10, 10), "视频文件", fill=(255, 255, 255), font=font)
            draw.text((10, 40), "点击播放", fill=(200, 200, 200), font=font)
            
            # 绘制播放按钮
            center_x, center_y = target_size[0] // 2, target_size[1] // 2
            button_size = 40
            draw.ellipse([center_x-button_size//2, center_y-button_size//2, 
                         center_x+button_size//2, center_y+button_size//2], 
                        fill=(255, 255, 255), outline=(0, 0, 0))
            
            # 绘制三角形播放图标
            triangle_points = [
                (center_x-10, center_y-15),
                (center_x-10, center_y+15),
                (center_x+15, center_y)
            ]
            draw.polygon(triangle_points, fill=(0, 0, 0))
            
            # 保存预览
            preview_path = self.get_preview_path(content_hash, size, 'jpg')
            img.save(preview_path, 'JPEG', quality=85, optimize=True)
            
            return {
                'success': True,
                'preview_path': str(preview_path),
                'preview_type': 'video',
                'size': img.size,
                'content_hash': content_hash
            }
            
        except Exception as e:
            logger.error(f"生成视频预览失败: {file_path}, 错误: {e}")
            return {
                'success': False,
                'error': f"视频预览生成失败: {str(e)}",
                'content_hash': content_hash
            }
    
    async def get_preview_info(self, content_hash: str) -> Dict[str, Any]:
        """获取预览信息"""
        try:
            preview_dir = self.cache_dir / content_hash
            
            if not preview_dir.exists():
                return {
                    'success': False,
                    'error': '预览不存在',
                    'content_hash': content_hash
                }
            
            # 获取所有预览文件
            preview_files = {}
            for size in self.preview_sizes.keys():
                for ext in ['jpg', 'png', 'gif']:
                    preview_path = preview_dir / f"{size}.{ext}"
                    if preview_path.exists():
                        preview_files[size] = {
                            'path': str(preview_path),
                            'size': preview_path.stat().st_size,
                            'modified': preview_path.stat().st_mtime
                        }
                        break
            
            return {
                'success': True,
                'content_hash': content_hash,
                'preview_files': preview_files
            }
            
        except Exception as e:
            logger.error(f"获取预览信息失败: {content_hash}, 错误: {e}")
            return {
                'success': False,
                'error': str(e),
                'content_hash': content_hash
            }
    
    async def cleanup_cache(self, max_age_days: int = 30) -> Dict[str, Any]:
        """清理缓存"""
        try:
            import time
            current_time = time.time()
            max_age_seconds = max_age_days * 24 * 60 * 60
            
            cleaned_count = 0
            total_size = 0
            
            for preview_dir in self.cache_dir.iterdir():
                if preview_dir.is_dir():
                    # 检查目录最后修改时间
                    if current_time - preview_dir.stat().st_mtime > max_age_seconds:
                        # 删除整个目录
                        import shutil
                        shutil.rmtree(preview_dir)
                        cleaned_count += 1
                    else:
                        # 计算目录大小
                        for file_path in preview_dir.rglob('*'):
                            if file_path.is_file():
                                total_size += file_path.stat().st_size
            
            return {
                'success': True,
                'cleaned_directories': cleaned_count,
                'total_size': total_size,
                'cache_dir': str(self.cache_dir)
            }
            
        except Exception as e:
            logger.error(f"清理缓存失败: {e}")
            return {
                'success': False,
                'error': str(e)
            }
